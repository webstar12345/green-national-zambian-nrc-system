"""
Report Export Service for NRC System
Handles PDF, Word, and Excel export functionality
"""

import io
from datetime import datetime
from django.http import HttpResponse
from django.utils import timezone
from django.conf import settings

# PDF Generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart

# Excel Generation
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.chart import PieChart, BarChart, Reference
import xlsxwriter

# Word Generation
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class ReportExporter:
    """Service class for exporting reports in various formats"""
    
    @staticmethod
    def export_summary_to_pdf(data, filename="summary_report.pdf"):
        """Export summary report to PDF"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.HexColor('#2D5016')  # Zambian Green
        )
        story.append(Paragraph("NRC System - Summary Report", title_style))
        
        # Generation info
        gen_info = f"Generated on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}"
        if data.get('date_from') or data.get('date_to'):
            date_range = f"Date Range: {data.get('date_from', 'Start')} to {data.get('date_to', 'End')}"
            gen_info += f"<br/>{date_range}"
        
        story.append(Paragraph(gen_info, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Statistics Table
        stats_data = [
            ['Metric', 'Count'],
            ['Total Applications', str(data['total_applications'])],
            ['Pending Applications', str(data['pending_count'])],
            ['Approved Applications', str(data['approved_count'])],
            ['Rejected Applications', str(data['rejected_count'])],
            ['New Applications', str(data['new_applications'])],
            ['Replacement Applications', str(data['replacement_applications'])],
            ['Male Applicants', str(data['male_count'])],
            ['Female Applicants', str(data['female_count'])],
            ['Recent Applications (30 days)', str(data['recent_applications'])],
        ]
        
        if data.get('avg_processing_time'):
            stats_data.append(['Average Processing Time (days)', str(data['avg_processing_time'])])
        
        stats_table = Table(stats_data, colWidths=[3*inch, 1.5*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2D5016')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
        ]))
        
        story.append(Paragraph("Application Statistics", styles['Heading2']))
        story.append(stats_table)
        story.append(Spacer(1, 20))
        
        # Top Districts Table
        if data.get('top_districts'):
            story.append(Paragraph("Top Districts by Application Volume", styles['Heading2']))
            districts_data = [['Rank', 'District', 'Applications', 'Percentage']]
            
            for i, district in enumerate(data['top_districts'][:10], 1):
                percentage = round((district['count'] / data['total_applications']) * 100, 1) if data['total_applications'] > 0 else 0
                districts_data.append([
                    str(i),
                    district['district'],
                    str(district['count']),
                    f"{percentage}%"
                ])
            
            districts_table = Table(districts_data, colWidths=[0.5*inch, 2*inch, 1*inch, 1*inch])
            districts_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D97706')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            
            story.append(districts_table)
        
        doc.build(story)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    @staticmethod
    def export_detailed_to_excel(applications, filename="detailed_report.xlsx"):
        """Export detailed report to Excel"""
        buffer = io.BytesIO()
        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "Applications Report"
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2D5016", end_color="2D5016", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        # Headers
        headers = [
            'Application ID', 'Applicant Name', 'Email', 'Application Type', 
            'Status', 'District', 'Village', 'Sex', 'Date of Birth',
            'NRC Number', 'Date Applied', 'Date Updated'
        ]
        
        for col, header in enumerate(headers, 1):
            cell = worksheet.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Data rows
        for row, app in enumerate(applications, 2):
            worksheet.cell(row=row, column=1, value=f"#{app.id:05d}")
            worksheet.cell(row=row, column=2, value=app.user.get_full_name())
            worksheet.cell(row=row, column=3, value=app.user.email)
            worksheet.cell(row=row, column=4, value=app.get_application_type_display())
            worksheet.cell(row=row, column=5, value=app.get_status_display())
            worksheet.cell(row=row, column=6, value=app.district)
            worksheet.cell(row=row, column=7, value=app.village)
            worksheet.cell(row=row, column=8, value='Male' if app.sex == 'M' else 'Female')
            worksheet.cell(row=row, column=9, value=app.date_of_birth.strftime('%Y-%m-%d'))
            worksheet.cell(row=row, column=10, value=app.nrc_number or 'Not Assigned')
            worksheet.cell(row=row, column=11, value=app.created_at.strftime('%Y-%m-%d %H:%M'))
            worksheet.cell(row=row, column=12, value=app.updated_at.strftime('%Y-%m-%d %H:%M') if app.updated_at else 'N/A')
        
        # Auto-adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # Add summary sheet
        summary_sheet = workbook.create_sheet("Summary")
        
        # Summary data
        total_apps = len(list(applications))
        pending = sum(1 for app in applications if app.status == 'pending')
        approved = sum(1 for app in applications if app.status == 'approved')
        rejected = sum(1 for app in applications if app.status == 'rejected')
        
        summary_data = [
            ['Report Summary', ''],
            ['Generated On', timezone.now().strftime('%Y-%m-%d %H:%M')],
            ['', ''],
            ['Total Applications', total_apps],
            ['Pending', pending],
            ['Approved', approved],
            ['Rejected', rejected],
        ]
        
        for row, (label, value) in enumerate(summary_data, 1):
            summary_sheet.cell(row=row, column=1, value=label).font = Font(bold=True)
            summary_sheet.cell(row=row, column=2, value=value)
        
        workbook.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    @staticmethod
    def export_summary_to_word(data, filename="summary_report.docx"):
        """Export summary report to Word document"""
        doc = Document()
        
        # Title
        title = doc.add_heading('NRC System - Summary Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generation info
        gen_info = f"Generated on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}"
        if data.get('date_from') or data.get('date_to'):
            gen_info += f"\nDate Range: {data.get('date_from', 'Start')} to {data.get('date_to', 'End')}"
        
        doc.add_paragraph(gen_info)
        doc.add_paragraph()  # Empty line
        
        # Statistics Section
        doc.add_heading('Application Statistics', level=1)
        
        # Create statistics table
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Count'
        
        # Data rows
        stats_data = [
            ('Total Applications', data['total_applications']),
            ('Pending Applications', data['pending_count']),
            ('Approved Applications', data['approved_count']),
            ('Rejected Applications', data['rejected_count']),
            ('New Applications', data['new_applications']),
            ('Replacement Applications', data['replacement_applications']),
            ('Male Applicants', data['male_count']),
            ('Female Applicants', data['female_count']),
            ('Recent Applications (30 days)', data['recent_applications']),
        ]
        
        if data.get('avg_processing_time'):
            stats_data.append(('Average Processing Time (days)', data['avg_processing_time']))
        
        for metric, count in stats_data:
            row_cells = stats_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(count)
        
        doc.add_paragraph()  # Empty line
        
        # Top Districts Section
        if data.get('top_districts'):
            doc.add_heading('Top Districts by Application Volume', level=1)
            
            districts_table = doc.add_table(rows=1, cols=4)
            districts_table.style = 'Table Grid'
            districts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            hdr_cells = districts_table.rows[0].cells
            hdr_cells[0].text = 'Rank'
            hdr_cells[1].text = 'District'
            hdr_cells[2].text = 'Applications'
            hdr_cells[3].text = 'Percentage'
            
            # Data rows
            for i, district in enumerate(data['top_districts'][:10], 1):
                percentage = round((district['count'] / data['total_applications']) * 100, 1) if data['total_applications'] > 0 else 0
                row_cells = districts_table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = district['district']
                row_cells[2].text = str(district['count'])
                row_cells[3].text = f"{percentage}%"
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    @staticmethod
    def export_exceptions_to_excel(exceptions, filename="exceptions_report.xlsx"):
        """Export exception report to Excel with formatting"""
        buffer = io.BytesIO()
        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "Exceptions Report"
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="DC2626", end_color="DC2626", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        # Headers
        headers = [
            'Application ID', 'Applicant Name', 'Issue Type', 'Severity',
            'Description', 'Days Pending', 'Status', 'Date Applied'
        ]
        
        for col, header in enumerate(headers, 1):
            cell = worksheet.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Data rows with conditional formatting
        for row, exc in enumerate(exceptions, 2):
            worksheet.cell(row=row, column=1, value=f"#{exc['application'].id:05d}")
            worksheet.cell(row=row, column=2, value=exc['application'].user.get_full_name())
            worksheet.cell(row=row, column=3, value=exc['issue_type'])
            
            # Severity cell with color coding
            severity_cell = worksheet.cell(row=row, column=4, value=exc['severity'])
            if exc['severity'] == 'Critical':
                severity_cell.fill = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")
                severity_cell.font = Font(color="DC2626", bold=True)
            elif exc['severity'] == 'High':
                severity_cell.fill = PatternFill(start_color="FED7AA", end_color="FED7AA", fill_type="solid")
                severity_cell.font = Font(color="EA580C", bold=True)
            else:
                severity_cell.fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
                severity_cell.font = Font(color="D97706", bold=True)
            
            worksheet.cell(row=row, column=5, value=exc['description'])
            worksheet.cell(row=row, column=6, value=exc['days_pending'])
            worksheet.cell(row=row, column=7, value=exc['application'].get_status_display())
            worksheet.cell(row=row, column=8, value=exc['application'].created_at.strftime('%Y-%m-%d'))
        
        # Auto-adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        workbook.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response